import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(16, 44, 87) # Navy Blue
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(192, 57, 43) # Deep Red Accent / Terracotta
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(41, 128, 185) # Steel Blue
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(f'"{text}"')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    # Process basic bold spans
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
        else:
            r = p.add_run(part)
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_bullet_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
        else:
            r = p.add_run(part)
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_table_data(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "102C57") # Deep Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(40, 40, 40)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

print("Helper library ready.")
