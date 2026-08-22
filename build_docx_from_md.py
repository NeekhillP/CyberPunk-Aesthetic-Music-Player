import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_helper import (
    add_styled_heading, add_quote, add_body_paragraph, 
    add_bullet_item, add_table_data
)

def build_word_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    
    # Title Cover Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("MOUNTAIN HELICOPTERS NEPAL")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 44, 87)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("Website Landing Page Sub-Tour Content Master File\nSEO-Optimized & Fact-Checked Copy for All 24 Helicopter Packages")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Let's read the markdown content source
    with open("c:/aesthetic music player/subtours_content_raw.md", "r", encoding="utf-8") as f:
        lines = f.readlines()

    current_table_headers = []
    current_table_rows = []
    in_table = False

    for line in lines:
        raw = line.strip()
        
        # Check if table line
        if raw.startswith("|") and raw.endswith("|"):
            cells = [c.strip() for c in raw.split("|")[1:-1]]
            # If separator row e.g. |:---|:---|
            if all(set(c).issubset({'-', ':', ' '}) for c in cells):
                continue
            if not in_table:
                in_table = True
                current_table_headers = cells
                current_table_rows = []
            else:
                current_table_rows.append(cells)
            continue
        else:
            if in_table:
                # Flush table
                add_table_data(doc, current_table_headers, current_table_rows)
                in_table = False
                current_table_headers = []
                current_table_rows = []

        if not raw:
            continue
            
        if raw.startswith("# CATEGORY") or raw.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(raw.replace("# ", "").replace("---", "—" * 25))
            r.font.name = 'Arial'
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(16, 44, 87)
            continue

        if raw.startswith("## ") and not raw.startswith("## H2"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(raw.replace("## ", ""))
            r.font.name = 'Arial'
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(192, 57, 43)
            continue

        # Format matchers
        if raw.startswith("H1:"):
            text = raw[3:].strip()
            add_styled_heading(doc, text, 1)
        elif raw.startswith("Quote:"):
            text = raw[6:].strip()
            add_quote(doc, text)
        elif raw.startswith("H2:"):
            text = raw[3:].strip()
            add_styled_heading(doc, text, 2)
        elif raw.startswith("H3:"):
            text = raw[3:].strip()
            add_styled_heading(doc, text, 3)
        elif raw.startswith("* ") or raw.startswith("- "):
            text = raw[2:].strip()
            add_bullet_item(doc, text)
        elif raw.startswith("Six Reasons to Book With Us"):
            add_styled_heading(doc, raw, 2)
        else:
            add_body_paragraph(doc, raw)

    if in_table:
        add_table_data(doc, current_table_headers, current_table_rows)

    output_path = "c:/aesthetic music player/Nepal_Helicopter_SubTours_Content.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    build_word_document()
